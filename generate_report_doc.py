#!/usr/bin/env python3
"""Generate analysis_report.docx with all charts embedded."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR  = os.path.join(BASE_DIR, "output")

doc = Document()

# ── Page setup ─────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Style config ────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# CJK font fallback
for s in doc.styles:
    rPr = s.element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.append(rFonts)

# Heading styles
for lvl, sz in [(1, 18), (2, 15), (3, 13)]:
    hs = doc.styles[f'Heading {lvl}']
    hs.font.size = Pt(sz)
    hs.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    rPr = hs.element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.append(rFonts)

# ── Helper functions ────────────────────────────────────────────

def add_img(name, width=Inches(5.8)):
    """Insert image from output folder."""
    path = os.path.join(OUT_DIR, name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(path, width=width)
        return True
    return False


def add_code(text):
    """Add code block styled paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_table(headers, rows):
    """Add formatted table."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Shading Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.bold = True
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            tbl.rows[i + 1].cells[j].text = str(val)
            for p in tbl.rows[i + 1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacer
    return tbl


# ══════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("TSAR 服务器监控数据分析报告")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("数据库设计 · 时间戳解析 · 小时聚合分析")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("数据范围: 20台服务器, 2026-07-01 ~ 2026-07-07 (7天)\n").font.size = Pt(11)
info.add_run("数据量: disk_tsar 12,000条 + pref_tsar 67,200条 = 79,200条\n").font.size = Pt(11)
info.add_run("工具: Python 3 + SQLite + matplotlib\n").font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  1. ER RELATIONSHIP DIAGRAM
# ══════════════════════════════════════════════════════════════

doc.add_heading("一、ER 实体关系图", level=1)

doc.add_paragraph(
    "下图展示了 4 张原始数据表的 ER 关系——采用左右分栏结构："
    "左侧为维度表（主机信息 + 指标字典），右侧为事实表（磁盘采集 + 性能采集），"
    "连线标注了关联字段和基数关系。"
)

add_img("01_er_diagram.png", width=Inches(5.5))

doc.add_heading("1.1 表间关联说明", level=2)

add_table(
    ["关联方向", "基数", "关联字段", "含义"],
    [
        ["HOST_DETAIL -> DISK_TSAR", "1 : N", "hostid", "一台主机产生多条磁盘采集记录 (每5分钟)"],
        ["HOST_DETAIL -> PREF_TSAR", "1 : N", "hostid", "一台主机产生多条性能采集记录 (每小时)"],
        ["MOD_DETAIL -> DISK_TSAR",  "1 : N", "mod",    "一个指标出现在多条磁盘记录中 (35个磁盘指标)"],
        ["MOD_DETAIL -> PREF_TSAR",  "1 : N", "mod",    "一个指标出现在多条性能记录中 (20个性能指标)"],
    ]
)

doc.add_heading("1.2 各表主外键设计", level=2)

add_table(
    ["表名", "主键", "外键", "记录数"],
    [
        ["HOST_DETAIL", "hostid (PK)", "—", "20"],
        ["MOD_DETAIL",  "mod (PK)",    "—", "55"],
        ["DISK_TSAR",   "ts + hostid + mod (联合唯一)", "hostid -> HOST_DETAIL, mod -> MOD_DETAIL", "12,000"],
        ["PREF_TSAR",   "ts + hostid + mod (联合唯一)", "hostid -> HOST_DETAIL, mod -> MOD_DETAIL", "67,200"],
    ]
)

p = doc.add_paragraph()
p.add_run("数据库设计说明：").font.bold = True
p.add_run("实际建库时将 DISK_TSAR 和 PREF_TSAR 合并为 tsar_detail 统一表，通过 type 字段 ('disk'/'pref') 区分类型；并另建 tsar_hourly 物化表存储小时汇总结果，查询更加方便。")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  2. TIMESTAMP PARSING
# ══════════════════════════════════════════════════════════════

doc.add_heading("二、时间戳解析", level=1)

doc.add_paragraph(
    "原始数据中 ts 字段是 13 位毫秒级 Unix 时间戳，"
    "表示从 1970-01-01 00:00:00 UTC 起经过的毫秒数。"
    "需要将其转换为北京时间 (UTC+8) 下的可读日期时间。"
)

add_img("02_timestamp_demo.png", width=Inches(5.5))

doc.add_heading("2.1 转换步骤", level=2)

add_table(
    ["步骤", "操作", "结果"],
    [
        ["1", "原始毫秒数", "1782835200000 (13位整数)"],
        ["2", "/ 1000 (毫秒 -> 秒)", "1782835200 (Unix timestamp)"],
        ["3", "fromtimestamp() -> UTC时间", "2026-06-30 16:00:00 UTC"],
        ["4", "+8小时 -> 北京时间", "2026-07-01 00:00:00 CST"],
    ]
)

doc.add_heading("2.2 Python 实现代码", level=2)

add_code("""from datetime import datetime, timezone, timedelta

BEIJING_TZ = timezone(timedelta(hours=8))

def ms_to_datetime(ts_ms: int) -> str:
    ts_sec = ts_ms / 1000.0        # step 1: ms -> seconds
    dt_utc = datetime.fromtimestamp(ts_sec, tz=timezone.utc)
    dt_beijing = dt_utc.astimezone(BEIJING_TZ)
    return dt_beijing.strftime("%Y-%m-%d %H:%M:%S")""")

doc.add_heading("2.3 转换验证", level=2)

add_table(
    ["原始毫秒数", "转换结果", "说明"],
    [
        ["1782835200000", "2026-07-01 00:00:00", "数据起点 (匹配文档)"],
        ["1782838800000", "2026-07-01 01:00:00", "1小时后"],
        ["1782921600000", "2026-07-02 00:00:00", "1天后"],
        ["1783436400000", "2026-07-07 23:00:00", "7天后 (pref终点)"],
    ]
)

p = doc.add_paragraph()
p.add_run("注意：").font.bold = True
p.add_run("数据库里同时保留了 ts_ms（原始毫秒值）和 dt（可读字符串）两列——原始值方便程序计算时间差，可读字符串方便人工查看和 SQL 分组。")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  3. HOURLY SUMMARIZATION
# ══════════════════════════════════════════════════════════════

doc.add_heading("三、按小时汇总指标", level=1)

doc.add_paragraph(
    "磁盘数据每 5 分钟采样一次，性能数据每小时采样一次。"
    "通过小时汇总，将同一小时内的多条记录聚合成一行，"
    "计算平均值 (AVG)、最大值 (MAX)、最小值 (MIN) 和采样数 (COUNT)。"
)

add_img("09_hourly_aggregation_demo.png", width=Inches(5.5))

doc.add_heading("3.1 汇总 SQL", level=2)

add_code("""CREATE VIEW v_tsar_hourly AS
SELECT
    strftime('%Y-%m-%d %H:00:00', dt) AS hour_bucket,
    hostid, type, mod, tag,
    ROUND(AVG(value), 4)  AS avg_value,     -- 小时均值
    ROUND(MAX(value), 4)  AS max_value,     -- 小时最大值
    ROUND(MIN(value), 4)  AS min_value,     -- 小时最小值
    COUNT(*)              AS sample_count   -- 采样数
FROM tsar_detail
GROUP BY hour_bucket, hostid, type, mod;""")

doc.add_heading("3.2 汇总统计", level=2)

add_table(
    ["指标", "数值"],
    [
        ["汇总总行数", "79,096"],
        ["小时时段数", "1,000"],
        ["覆盖主机", "20 / 20 (100%)"],
        ["覆盖指标", "55 / 55 (100%)"],
        ["磁盘类每小时采样数", "平均 ~12 条 (5分钟间隔)"],
        ["性能类每小时采样数", "1 条 (整点采样)"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  4. DATA ANALYSIS & VISUALIZATIONS
# ══════════════════════════════════════════════════════════════

doc.add_heading("四、数据分析与可视化", level=1)

# 4.1 CPU
doc.add_heading("4.1 CPU 使用率分析", level=2)
add_img("03_cpu_hourly_trend.png", width=Inches(5.5))
add_img("04_cpu_host_comparison.png", width=Inches(5.5))

doc.add_paragraph("分析发现：")
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("全局平均 CPU：").font.bold = True
p.add_run("43.2%，各主机平均分布在 39%~48% 之间，差异不大。")

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("峰值 CPU：").font.bold = True
p.add_run("普遍在 90% 以上，部分主机达到 99%，说明存在周期性负载高峰。")

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("趋势判断：").font.bold = True
p.add_run("CPU 呈现规律波动，无异常突增或持续满载，整体负载健康。")

# 4.2 Disk
doc.add_heading("4.2 磁盘使用率分析", level=2)
add_img("06_disk_util_trend.png", width=Inches(5.5))

doc.add_paragraph("分析发现：")
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("Top 5 峰值磁盘主机：").font.bold = True
p.add_run("host010 (99.5%), host020 (99.0%), host003 (99.0%), host017 (98.9%), host014 (97.0%)。")

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("磁盘 sda 使用率峰值接近 100%，建议关注这些主机是否需要扩容或优化 I/O。")

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("监控覆盖 5 块磁盘 (sda~sde)，每块盘 7 种指标：rqm、read、write、avgrq、await、util、svctm。")

# 4.3 Memory
doc.add_heading("4.3 内存使用分析", level=2)
add_img("05_memory_overview.png", width=Inches(5.5))

doc.add_paragraph("分析发现：")
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("mem_used 稳定在较高水平（host001 约 80~90 GB），内存占用较大。")

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("mem_free + mem_buff + mem_cache 构成可用内存缓冲。")

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("mem_swap 使用量很低，说明物理内存充足，未出现严重的内存压力。")

# 4.4 Network
doc.add_heading("4.4 网络流量分析", level=2)
add_img("07_network_traffic.png", width=Inches(5.5))

doc.add_paragraph("分析发现：")
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("入站 (net_in) 和出站 (net_out) 带宽呈现互补波动，符合请求-响应通信模式。")

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("数据包速率 (net_pktin / net_pktout) 波动较大，存在明显的高峰和低谷时段。")

# 4.5 Load
doc.add_heading("4.5 系统负载分析", level=2)
add_img("08_load_average.png", width=Inches(5.5))

doc.add_paragraph("分析发现：")
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("load1 < load5 < load15 关系偶尔打破，短期负载有尖峰。")

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("load1 波动剧烈、load15 相对平滑，符合负载平均的数学特性。")

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("整体负载在合理范围，未发现持续高负载。")

# 4.6 Dashboard
doc.add_heading("4.6 综合仪表盘", level=2)
add_img("10_dashboard.png", width=Inches(5.5))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  5. SUMMARY
# ══════════════════════════════════════════════════════════════

doc.add_heading("五、总结", level=1)

doc.add_heading("5.1 数据库设计", level=2)

doc.add_paragraph(
    "采用星型模型：tsar_detail 作为事实表，host_detail 和 mod_detail 作为维度表，"
    "tsar_hourly 作为预聚合表。"
)

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("查询高效：").font.bold = True
p.add_run("维度表小，事实表有针对性的索引。")

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("存储合理：").font.bold = True
p.add_run("同时保留原始毫秒时间戳和可读日期时间，兼顾精度和易用性。")

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run("扩展方便：").font.bold = True
p.add_run("新增主机或指标只需插入维度表，不影响现有结构。")

doc.add_heading("5.2 核心分析结论", level=2)

add_table(
    ["维度", "结论", "详情"],
    [
        ["CPU",  "整体健康", "平均 43.2%，峰值偶达 99%，属正常波动"],
        ["磁盘", "需关注",   "5台主机 sda 峰值超过 97%，建议扩容或优化 I/O"],
        ["内存", "充足",     "swap 使用极低，物理内存配置合理"],
        ["网络", "正常",     "流量波动符合业务规律，无异常流量"],
        ["负载", "可控",     "整体负载在合理范围，未发现持续高负载"],
    ]
)

doc.add_heading("5.3 提交文件清单", level=2)

add_table(
    ["分类", "文件", "说明"],
    [
        ["ER图",      "01_er_diagram.png",          "ER 实体关系图 (左右分栏, 1:N箭头)"],
        ["时间戳",    "02_timestamp_demo.png",       "时间戳转换演示 (5个示例)"],
        ["小时汇总",  "09_hourly_aggregation_demo.png", "5分钟原始 -> 小时汇总对比"],
        ["CPU分析",   "03_cpu_hourly_trend.png, 04_cpu_host_comparison.png", "趋势 + 主机对比"],
        ["内存分析",  "05_memory_overview.png",     "内存各组件趋势"],
        ["磁盘分析",  "06_disk_util_trend.png",     "磁盘使用率趋势"],
        ["网络分析",  "07_network_traffic.png",     "网络带宽 + 数据包"],
        ["负载分析",  "08_load_average.png",        "load1/5/15 趋势"],
        ["仪表盘",   "10_dashboard.png",            "四合一综合概览"],
        ["建表SQL",  "create_schema.sql",           "完整建表语句 (3表+2视图+索引)"],
        ["ETL脚本",  "etl_import.py",               "数据导入 + 时间戳转换"],
        ["汇总脚本", "hourly_summary.py",           "小时汇总 + CSV 导出"],
        ["数据库",   "tsar_monitor.db",             "SQLite 数据库文件"],
    ]
)

# ── Save ───────────────────────────────────────────────────────

out_path = os.path.join(BASE_DIR, "analysis_report.docx")
doc.save(out_path)
print(f"[OK] Report saved to: {out_path}")
